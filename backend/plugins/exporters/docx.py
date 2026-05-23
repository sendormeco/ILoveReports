import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from backend.domain.models import DocxRequest
from backend.plugins.base import AbstractExporter, ExportedDocument


class DocxExporter(AbstractExporter):
    """Exporter plugin that creates a Microsoft Word .docx document."""

    @property
    def name(self) -> str:
        return "docx"

    @property
    def file_extension(self) -> str:
        return ".docx"

    @property
    def media_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self, payload: DocxRequest) -> ExportedDocument:
        document = Document()

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(f"{payload.title}\n\n")
        title_run.bold = True
        title_run.font.size = Pt(24)

        self._add_centered_line(document, f"Автор: {payload.author}")
        self._add_centered_line(document, f"Организация: {payload.organization}")
        if payload.supervisor:
            self._add_centered_line(document, f"Руководитель: {payload.supervisor}")
        self._add_centered_line(document, f"Дата: {payload.date}")

        document.add_page_break()

        for line in payload.report.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            paragraph = document.add_paragraph()
            if stripped.startswith("# "):
                run = paragraph.add_run(stripped[2:])
                run.bold = True
                run.font.size = Pt(18)
            elif stripped.startswith("## "):
                run = paragraph.add_run(stripped[3:])
                run.bold = True
                run.font.size = Pt(15)
            else:
                paragraph.add_run(stripped)

        buffer = io.BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        return ExportedDocument(
            content=content,
            filename=self.safe_download_name(payload.title),
            media_type=self.media_type,
            file_extension=self.file_extension,
        )

    @staticmethod
    def safe_download_name(title: str) -> str:
        return AbstractExporter.safe_filename(title, ".docx")

    @staticmethod
    def _add_centered_line(document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{text}\n")
        run.font.size = Pt(14)
