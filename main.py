from fastapi import FastAPI
from pydantic import BaseModel, Extra
from fastapi.middleware.cors import CORSMiddleware
import subprocess
from fastapi.responses import FileResponse
from docx import Document
import uuid
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


app = FastAPI(title="ILoveReports MVP - Ollama LLaMA3")

class DocxRequest(BaseModel):
    title: str
    report: str
    author: str
    organization: str
    date: str
    supervisor: str | None = None

    class Config:
        extra = Extra.allow


class ReportRequest(BaseModel):
    goal: str
    process: str
    results: str
    conclusion: str
    tone: str | None = None

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # разрешённые источники
    allow_credentials=True,
    allow_methods=["*"],     # разрешаем все методы (GET, POST и т.д.)
    allow_headers=["*"],     # разрешаем любые заголовки
)

@app.post("/generate-docx")
def generate_docx(payload: DocxRequest):
    doc = Document()

    # ======= Титульный лист =======
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(payload.title + "\n\n")
    title_run.bold = True
    title_run.font.size = Pt(24)

    # Автор
    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_paragraph.add_run(f"Автор: {payload.author}\n")
    author_run.font.size = Pt(14)

    # Организация
    org_paragraph = doc.add_paragraph()
    org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_paragraph.add_run(f"Организация: {payload.organization}\n")
    org_run.font.size = Pt(14)

    # Руководитель (если есть)
    if payload.supervisor:
        sup_paragraph = doc.add_paragraph()
        sup_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sup_run = sup_paragraph.add_run(f"Руководитель: {payload.supervisor}\n")
        sup_run.font.size = Pt(14)

    # Дата
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Дата: {payload.date}\n")
    date_run.font.size = Pt(14)

    # Разрыв страницы после титульного листа
    doc.add_page_break()

    # ======= Основной текст отчёта =======
    for line in payload.report.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    # ======= Сохранение файла =======
    filename = f"report_{uuid.uuid4()}.docx"
    temp_dir = os.path.join(os.getcwd(), "temp_docs")
    os.makedirs(temp_dir, exist_ok=True)
    path = os.path.join(temp_dir, filename)

    doc.save(path)

    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{payload.title}.docx"
    )


@app.post("/generate")
async def generate_report(data: ReportRequest):
    prompt = (
        "Ты — помощник для написания лабораторных отчётов на русском языке. "
        "Составь развёрнутый, подробный и длинный отчёт по предоставленным данным. "
        "Не используй английский язык, не подставляй свои примеры или цифры. "
        f"Используй данные: Цель: {data.goal}; Ход работы: {data.process}; "
        f"Результаты: {data.results}; Вывод: {data.conclusion}.\n"
    )

    if data.tone:
        prompt += f"Отчёт должен быть написан с учётом следующего пожелания по стилю и тону: {data.tone}\n"

    prompt += (
        "Сделай текст связным, в абзацах, с разделами: Введение, Цель работы, Ход работы, Результаты, Вывод. "
        "Пиши только по фактам."
    )

    try:
        # Запуск ollama run llama3
        process = subprocess.Popen(
            [
                r"C:\Users\johndoe\AppData\Local\Programs\Ollama\ollama.exe",
                "run",
                "llama3"
            ],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,         # важный момент для текста
            encoding="utf-8"   # чтобы не падало на русских символах
        )

        # Передаем промпт в stdin и получаем stdout, stderr
        stdout, stderr = process.communicate(prompt, timeout=300)

        if process.returncode != 0:
            # Если ollama завершился с ошибкой
            return {"error": f"Ollama failed: {stderr.strip()}"}

        # Возвращаем отчёт, обрезаем лишние пробелы
        report_text = stdout.strip()
        if not report_text:
            report_text = "Ollama вернул пустой ответ."

        return {"report": report_text}

    except subprocess.TimeoutExpired:
        process.kill()
        return {"error": "Ollama timed out"}

    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"}

@app.get("/")
def root():
    return {"status": "ILoveReports backend with Ollama LLaMA3 is running"}
